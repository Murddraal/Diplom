import requests
import json
import csv
from docx import Document

class translator(object):
    def __init__(self, lang = 'en'):
        self.url = 'https://translate.yandex.net/api/v1.5/tr.json/translate?'
        self.key = 'trnsl.1.1.20170127T090047Z.d0549b75a2237806.40d7e9d8902cb56bc4115af8e88c1bf8f0e4cbf3'
        self.lang = lang


    def csv_reading(self, csv_filename):
        with open(csv_filename) as csvfile:
            reader = csv.reader(csvfile, dialect='excel-tab')
            for row in reader:
                yield row

    def writing(self, file_name, orig_text, trans_text):
        document = Document()
        paragraph = document.add_paragraph('New Paragraph')
        document.add_heading('HEAD')
        table = document.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = 'First cell'
        row.cells[1].text = 'Second cell'
        document.save('test.doc')

    def text_parsing(self, json_row):
        mail = json.loads(json_row)

        return mail['bodyText']


    def translating(self, text):
        translation = requests.post(self.url, data={'key': self.key, 'text': text, 'lang': self.lang})

        jtr = json.loads(translation.text)

        return jtr['text'][0]