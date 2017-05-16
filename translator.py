"""This module translate text, detect languages.
"""
import sys
import csv
import json
import operator
import re
import requests
from docx import Document
from docx.opc import exceptions as doc_exc
from docx.shared import Pt
from langdetect import detect_langs
from langdetect import detect
from langdetect import lang_detect_exception as ld_exc
# coding: utf8

# DetectorFactory.seed = 0


class Translator(object):
    """This class can translate text via yandex.translate service,
    detect language via langdetect and yandex translate.
    """

    def __init__(self, lang='en'):
        self.url_tr = 'https://translate.yandex.net/api/v1.5/tr.json/translate?'
        self.url_det = 'https://translate.yandex.net/api/v1.5/tr.json/detect? '
        self.key = 'trnsl.1.1.20161229T092652Z.92eec6478cc791c5.9a428e7546269ce12b04df552805c48a4c63261e'
        self.lang = lang
        self.max_text_size = 10000
        self.languages = {}

    def translating(self, text):
        """translate text via yandex.translate
        """
        splited_text = re.split(r'[?!:\n\.]', text)

        translated = ""
        for i in splited_text:
            if len(i) > self.max_text_size:
                return translated
            if len(re.findall(r'[^\s?1\w\n]', i)) == 0:
                continue
            translation = requests.post(self.url_tr,
                                        data={'key': self.key,
                                              'text': i,
                                              'lang': self.lang})

            jtr = json.loads(translation.text)
            translated += jtr['text'][0] + " ||| "
        return translated

    def lang_detecting(self, text):
        """detect language via langdetect
        """

        try:
            lg = detect_langs(text)
        except TypeError:
            return
        except ld_exc.LangDetectException:
            return

        list_lg = [str(i).split(':') for i in lg]

        for i in list_lg:
            try:
                self.languages[i[0]] += 1
            except KeyError:
                self.languages[i[0]] = 1

    def print_langs_list(self, f_name):
        """print all detected language with it's weight in file
        """
        with open(f_name, 'w') as file:
            for i in self.languages:
                file.write(str(i[0]) + '-' + str(i[1]) + '\n')


def csv_reading(csv_filename):
    """It's a generator.
    Get path to the .csv file and return it's content line by line
    """

    with open(csv_filename, encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile, dialect='excel-tab')
        for row in reader:
            str_row = ""
            for cell in row:
                str_row += cell
            yield str_row


def writing(f_name, text):
    """Write text in file
    """
    if f_name.endswith(".doc"):
        try:
            document = Document(f_name)
        except doc_exc.PackageNotFoundError:
            document = Document()

        run = document.add_paragraph().add_run()
        font = run.font
        font.name, font.size = 'Times New Roman', Pt(6)
        table = document.add_table(rows=1, cols=1)
        row = table.rows[0]
        row.cells[0].text = text
        document.save(f_name)

    elif f_name.endswith(".txt"):
        with open(f_name, 'w') as file:
            file.write(text)

def writing_langs(f_name, langs):
    with open(f_name, 'w') as file:
        for l in langs:
            file.write("{} - {}\n".format(l[0], l[1]))


def text_parsing_json(json_row, attr):
    """Parsing text from json. The text takes from attribute attr
    """
    mail = json.loads(json_row)
    try:
        value = mail[attr]
    except ValueError:
        print("There is no attr like that")
        exit(-1)
    return value
