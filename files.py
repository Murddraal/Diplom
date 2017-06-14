from docx import Document
from docx.opc import exceptions as doc_exc
from docx.shared import Pt 
import csv
import json

def csv_reading(csv_filename, in_one_row=True, dialect='excel-tab'):
    """It's a generator.
    Get path to the .csv file and return it's content line by line
    """

    with open(csv_filename, encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile
                            ,dialect=dialect
                            )
        for row in reader:
            str_row = ""
            if in_one_row:
                for cell in row:
                    str_row += cell
                yield str_row
            else:
                yield row


def writing(f_name, text):
    """Write text in file
    """
    if f_name.endswith(".doc"):
        try:
            document = Document(f_name)
        except doc_exc.PackageNotFoundError:
            document = Document()

        run = document.add_paragraph().add_run()
        font = run.font
        font.name, font.size = 'Times New Roman', Pt(6)
        table = document.add_table(rows=1, cols=1)
        row = table.rows[0]
        row.cells[0].text = text
        document.save(f_name)

    elif f_name.endswith(".txt"):
        with open(f_name, 'w') as file:
            file.write(text)

def writing_langs(f_name, langs):
    with open(f_name, 'w') as file:
        for l in langs:
            file.write("{} - {}\n".format(l[0], l[1]))


def text_parsing_json(json_row, attr):
    """Parsing text from json. The text takes from attribute attr
    """
    mail = json.loads(json_row)
    try:
        value = mail[attr]
    except ValueError:
        print("There is no attr like that")
        exit(-1)
    return value